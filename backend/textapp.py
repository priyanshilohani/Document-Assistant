from flask import Flask, request, jsonify
from sentence_transformers import SentenceTransformer
from transformers import BartForConditionalGeneration, BartTokenizer
import pdfplumber
import docx 
import re
import spacy
import numpy as np
import jwt
import datetime
import logging
import os
from pymongo import MongoClient
from flask_cors import CORS
from functools import wraps
from bson import ObjectId
import PyPDF2
from docx import Document

app = Flask(__name__)
CORS(app, supports_credentials=True, origins=["http://localhost:3000"])
logging.basicConfig(level=logging.INFO)

SECRET_KEY = os.getenv("JWT_SECRET_KEY", "default_secret_key")

# Connecting to MongoDB
try:
    mongo_client = MongoClient("mongodb://localhost:27017/")
    db = mongo_client["DocAssist"]
    documents_collection = db["documents"]
    logging.info("Connected to MongoDB successfully!")
except Exception as e:
    logging.error(f"Error connecting to MongoDB: {e}")
    exit(1)

# Loading models
try:
    bart_model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn").to("cpu")
    bert_model = SentenceTransformer('all-MiniLM-L6-v2', device="cpu")
    bart_tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    logging.error(f"Error loading models: {e}")
    exit(1)

# Authentication 
def auth_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get("Authorization")
        if token:
            token_parts = token.split(" ")
            if len(token_parts) == 2 and token_parts[0] == "Bearer":
                token = token_parts[1]
            else:
                return jsonify({"error": "Invalid Authorization format!"}), 401
        else:
            token = request.cookies.get("token")
            if not token:
                return jsonify({"error": "Authentication token is missing!"}), 401
        try:
            decoded_token = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
            request.user_id = decoded_token.get("user_id")
            if not request.user_id:
                return jsonify({"error": "Invalid token: user_id missing!"}), 401
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token has expired!"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Invalid token!"}), 401
        return f(*args, **kwargs)
    return decorated_function

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)
def extract_text_from_file(file):
    if file.filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file.filename.endswith(".docx"):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return file.read().decode("utf-8")
def save_document_to_db(title, content, user_id=None):
    documents_collection.insert_one({
        "title": title,
        "content": content,
        "user_id": user_id,
    })
import random
import string
def generate_document_id(length=24):
    characters = string.ascii_letters + string.digits
    return ''.join(random.choice(characters) for _ in range(length))
def get_user_id_from_token():
    token = None
    # First try header
    auth_header = request.headers.get("Authorization")
    if auth_header and auth_header.startswith("Bearer "):
        token = auth_header.split(" ")[1]
    else:
        token = request.cookies.get("token")  # fallback to cookie
    if not token:
        return None, {"error": "Token missing!"}, 401
    try:
        print(f"Incoming token: {token}")
        # Inspect payload without verifying to debug
        unverified_payload = jwt.decode(token, options={"verify_signature": False})
        print(f"nverified Payload: {unverified_payload}")
        # Now try to decode properly
        decoded_token = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        print(f"Token successfully verified")
        user_id = decoded_token.get("user_id")
        if not user_id:
            return None, {"error": "Invalid token: user_id missing!"}, 401
        return user_id, None, 200
    except jwt.ExpiredSignatureError:
        print("Token expired")
        return None, {"error": "Token has expired!"}, 401
    except jwt.InvalidTokenError as e:
        print(f"Invalid Token Error: {e}")
        return None, {"error": "Invalid token!"}, 401

@app.route('/save-progress', methods=['POST'])
@auth_required
def save_progress():
    try:
        user_id, error_response, status_code = get_user_id_from_token()
        if not user_id:
            return jsonify(error_response), status_code
        data = request.get_json()
        content = data.get('content')
        title = data.get('title')
        if not content or not title:
            return jsonify({"error": "Missing title or content"}), 400
        cleaned_content = clean_text(content)
        if not cleaned_content:
            return jsonify({"error": "Content is empty after cleaning"}), 400
        chunks = split_into_chunks(cleaned_content)
        embeddings = [{"text": chunk, "embedding": bert_model.encode(chunk).tolist()} for chunk in chunks]
        if not embeddings:
            return jsonify({"error": "No valid embeddings generated"}), 400
        document_id = generate_document_id()
        document_entry = {
            "user_id": user_id,
            "document_id": document_id,
            "document_name": f"{title}.txt",
            "document": cleaned_content,
            "embeddings": embeddings,
            "timestamp": datetime.datetime.utcnow()
        }
        documents_collection.insert_one(document_entry)
        logging.info(f"Progress saved for user: {user_id}")
        return jsonify({"message": "Progress saved successfully", "document_id": document_id}), 200
    except Exception as e:
        logging.error(f"Error saving progress: {str(e)}")
        return jsonify({"error": f"Error saving progress: {str(e)}"}), 500

@app.route("/open-document", methods=["POST"])
def open_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    text = extract_text_from_file(file)
    return jsonify({"text": text})

def clean_text(text):
    return re.sub(r"\s+", " ", text).strip()

def split_into_chunks(text, max_chunk_size=512):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    chunks, current_chunk = [], ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chunk_size:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def cosine_similarity(vec1, vec2):
    return np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))

@app.route("/")
def home():
    return "Welcome to the DocAssist API!"

@app.route("/process-document", methods=["POST"])
@auth_required
def process_document():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file provided!"}), 400
        logging.info(f"Processing file: {file.filename}")
        if file.filename.endswith(".txt"):
            content = file.read().decode("utf-8")
        elif file.filename.endswith(".pdf"):
            content = extract_text_from_pdf(file)
        elif file.filename.endswith(".docx"):
            content = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format!"}), 400
        cleaned_content = clean_text(content)
        if not cleaned_content:
            return jsonify({"error": "Extracted content is empty!"}), 400
        chunks = split_into_chunks(cleaned_content)
        embeddings = [{"text": chunk, "embedding": bert_model.encode(chunk).tolist()} for chunk in chunks]
        if not embeddings:
            return jsonify({"error": "No valid embeddings generated!"}), 400
        document_id = str(ObjectId())
        document_entry = {
            "user_id": request.user_id,
            "document_id": document_id,
            "document_name": file.filename,
            "document": cleaned_content,
            "embeddings": embeddings,
            "timestamp": datetime.datetime.utcnow()
        }
        documents_collection.insert_one(document_entry)
        logging.info("Document successfully saved in MongoDB")
        return jsonify({"message": "File processed and embedded successfully!", "document_id": document_id})
    except Exception as e:
        logging.error(f"Error processing document: {str(e)}")
        return jsonify({"error": f"Error processing document: {str(e)}"}), 500
    
@app.route("/get-uploads", methods=["GET"])
def get_uploaded_files():
    try:
        files = os.listdir("uploads")  # Assuming files are stored in the 'uploads' directory
        return jsonify({"files": files})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/get-suggestions", methods=["POST"])
@auth_required
def get_suggestions():
    try:
        data = request.json
        query = data.get("query", "")
        if not query:
            logging.info("Received empty query")
            return jsonify({"error": "Query is required!"}), 400
        logging.info(f"Received query: {query}")
        user_documents = list(documents_collection.find({"user_id": request.user_id}))
        if not user_documents:
            logging.info(f"No documents found for user: {request.user_id}")
            return jsonify({"error": "No documents found for the user!"}), 404
        logging.info(f"Found {len(user_documents)} documents for user.")
        query_embedding = bert_model.encode(query)
        logging.info(f"Query embedding shape: {np.array(query_embedding).shape}")
        ranked_chunks = []
        seen_snippets = set()  # To track and remove duplicates
        for doc in user_documents:
            doc_name = doc.get("document_name", "Unknown")
            logging.info(f"Processing document: {doc_name}")
            for chunk in doc.get("embeddings", []):
                text_snippet = chunk["text"].strip()
                if text_snippet in seen_snippets:
                    continue  
                similarity = cosine_similarity(query_embedding, chunk["embedding"])
                logging.info(f"Comparing chunk: '{text_snippet[:50]}...' | Similarity: {similarity:.4f}")
                # similarity threshold 
                if similarity > 0.6:
                    ranked_chunks.append({
                        "filename": doc_name,
                        "content_snippet": text_snippet,
                        "similarity": similarity
                    })
                    seen_snippets.add(text_snippet)  # Mark this snippet as seen
        # Sort by similarity in descending order
        ranked_chunks.sort(key=lambda x: x["similarity"], reverse=True)
        max_results = min(5, len(ranked_chunks))
        min_results = min(3, len(ranked_chunks))
        filtered_chunks = ranked_chunks[:max_results] if max_results > 0 else ranked_chunks[:min_results]
        logging.info(f"Final suggestions count: {len(filtered_chunks)}")
        if not filtered_chunks:
            logging.info("No chunks passed similarity threshold")
            return jsonify({"error": "No relevant information found!"}), 404
        return jsonify({"suggestions": filtered_chunks}), 200
    except Exception as e:
        logging.error(f"Error generating suggestions: {str(e)}")
        return jsonify({"error": f"Error generating suggestions: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5002)

